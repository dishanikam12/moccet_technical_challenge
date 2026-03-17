#!/usr/bin/env python3
"""
Convert docs/*.md to docs/*.docx.
Run from project root: python scripts/docs_to_docx.py
"""
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"


def parse_table_line(line: str) -> list[str]:
    """Split a markdown table row into cells (strip whitespace)."""
    parts = line.split("|")
    if len(parts) < 2:
        return []
    return [c.strip() for c in parts[1:-1]]


def is_separator_line(line: str) -> bool:
    """True if line is like |------|------|."""
    line = line.strip()
    if not line or not line.startswith("|"):
        return False
    return re.match(r"^\|[\s\-:]+\|", line) is not None


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    table_rows: list[list[str]] = []
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        raw = line
        line = line.rstrip()

        # Code block
        if line.strip().startswith("```"):
            if in_code_block:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Normal"
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            if is_separator_line(line):
                i += 1
                continue
            cells = parse_table_line(line)
            if cells:
                table_rows.append(cells)
            i += 1
            continue
        else:
            # Flush table if any
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=ncols)
                t.style = "Table Grid"
                for ri, row_cells in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < ncols:
                            t.rows[ri].cells[ci].text = cell_text
                table_rows = []

        # List item
        if line.strip().startswith("- ") or re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(line.strip(), style="List Bullet")
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        doc.add_paragraph(line.strip())
        i += 1

    # Flush remaining table
    if table_rows:
        ncols = max(len(r) for r in table_rows)
        t = doc.add_table(rows=len(table_rows), cols=ncols)
        t.style = "Table Grid"
        for ri, row_cells in enumerate(table_rows):
            for ci, cell_text in enumerate(row_cells):
                if ci < ncols:
                    t.rows[ri].cells[ci].text = cell_text

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def main():
    try:
        from docx import Document
    except ImportError:
        print("Install python-docx: pip install python-docx")
        return 1
    for name in ["llm_judge_cost", "llm_judge_prompt"]:
        md_path = DOCS / f"{name}.md"
        if not md_path.exists():
            print(f"Skip (not found): {md_path}")
            continue
        docx_path = DOCS / f"{name}.docx"
        md_to_docx(md_path, docx_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
